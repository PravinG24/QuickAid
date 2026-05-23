from docx import Document
from datetime import datetime
from pathlib import Path
import shutil

src = Path(r"D:\Documents\MyMahir Docs\Course Work\Project Submissions\QuickAid_WriteUp.docx")
if not src.exists():
    print('Write-up not found:', src)
    raise SystemExit(1)

# create timestamped backup
bak = src.with_name(src.stem + '_backup_' + datetime.now().strftime('%Y%m%d_%H%M%S') + src.suffix)
shutil.copy2(src, bak)
print('Backup created:', bak)

# open existing document
doc = Document(str(src))

# Append an implementation update section with code-level details
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(text):
    doc.add_paragraph(text)

def add_bullets(items):
    for it in items:
        doc.add_paragraph('- ' + it)

add_heading('Implementation Update — Code-level Details', level=1)
add_paragraph('This section records concrete, code-level behaviours discovered during code review and test of the QuickAid backend and frontend.')

add_heading('Secrets and Key Vault behaviour', level=2)
add_bullets([
    'Secret loader (`shared/secrets.py`): attempts Key Vault lookup using DefaultAzureCredential when `KEY_VAULT_URL` is set, otherwise reads an environment variable provided by the caller (`env_fallback`). Raises `RuntimeError` if the secret cannot be resolved.',
    'Functions run with Managed Identity in Azure; locally DefaultAzureCredential falls back to developer credentials when available.',
])

add_heading('Activity logs and notifications', level=2)
add_bullets([
    'Activity logs are persisted to a blob-backed JSON store via `shared/blob_store.py` and built with `create_activity_log(...)` in `shared/activity_log.py`.',
    'Activity log fields: `id` (LOG-<timestamp>-<ticket_id>), `actor`, `actor_type`, `action`, `ticket_id`, `timestamp`, `updated_fields`, `old_values`.',
    'Notifications are persisted similarly via `shared/notifications.py` using `create_notification(...)`.',
    'Notification fields: `id` (NOTIF-<uuid4>), `email`, `ticket_id`, `message`, `updated_fields`, `timestamp`, `read` (bool).',
])

add_heading('Implemented HTTP endpoints (summary)', level=2)
add_bullets([
    'POST `/api/submit_ticket` — create ticket; writes to Cosmos DB; triggers SendGrid confirmation (non-blocking on failure).',
    'GET `/api/get_ticket` and GET `/api/tickets` — retrieve ticket(s) by `email` or `ticketId`.',
    'POST `/api/register_user`, POST `/api/user_login` — user registration and authentication flows.',
    'POST `/api/register_admin` and GET/PATCH `/api/admin_approvals` — admin requests and approval workflow.',
    'PATCH `/api/tickets_update/{ticketId}` and POST `/api/bulk_update_tickets` — admin-only updates that create activity logs and notifications for tracked fields (status, priority, assignedTo, category).',
    'GET `/api/notifications` and PATCH `/api/notifications/{notificationId}` — retrieve and mark notifications; `notificationId=all` supported with `email` param.',
    'GET `/api/get_activity_log` — returns activity log entries for a ticket (from blob JSON).',
])

add_heading('Admin behaviour and side-effects', level=2)
add_bullets([
    'Admin update endpoints validate admin role via `shared/admin_auth.py` and `shared/jwt_utils.py`.',
    'On tracked field changes, handlers call `create_activity_log(...)` and `create_notification(...)` per affected ticket.',
    'Bulk updates iterate per-ticket, applying same logging and notification behaviour for each ticket.',
])

add_heading('Representative function signatures', level=2)
add_bullets([
    '`create_activity_log(actor, actor_type, action, ticket_id, updated_fields, old_values)` — builds and persists a log entry.',
    '`create_notification(email, ticket_id, message, updated_fields)` — creates and persists a notification; `read` defaults to `False`.',
    '`get_secret(secret_name, env_fallback=None)` — returns secret value; Key Vault first then env fallback, else throws.',
])

add_heading('Notes on persistence choices', level=2)
add_paragraph('Blob-backed JSON store is simple and reliable for prototype scope; recommend migrating logs/notifications to a database (Cosmos DB or Table Storage) for production-scale concurrency and querying.')

# Save updated document to a new file (avoid write-lock issues with Word)
new_path = src.with_name(src.stem + '_updated' + src.suffix)
try:
    doc.save(str(new_path))
    print('Updated write-up saved to', new_path)
except Exception as e:
    print('Failed to save updated write-up:', e)
    raise

# Basic validation: confirm the new file contains the inserted heading
doc2 = Document(str(new_path))
found = any('Implementation Update' in p.text for p in doc2.paragraphs)
if found:
    print('Validation: inserted heading found in', new_path)
else:
    print('Validation: inserted heading NOT found in', new_path)
