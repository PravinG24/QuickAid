from docx import Document
from datetime import datetime
from pathlib import Path
import shutil
from docx.shared import Pt

src = Path(r"D:\Documents\MyMahir Docs\Course Work\Project Submissions\QuickAid_WriteUp_updated.docx")
if not src.exists():
    print('Updated write-up not found:', src)
    raise SystemExit(1)

# backup
bak = src.with_name(src.stem + '_refine_backup_' + datetime.now().strftime('%Y%m%d_%H%M%S') + src.suffix)
shutil.copy2(src, bak)
print('Backup of updated file created:', bak)

# open
doc = Document(str(src))

def add_h(level, text):
    doc.add_heading(text, level=level)

def add_p(text):
    doc.add_paragraph(text)

def add_code_block(code_lines):
    p = doc.add_paragraph()
    for i, line in enumerate(code_lines):
        run = p.add_run(line + ("\n" if i < len(code_lines)-1 else ""))
        font = run.font
        font.name = 'Consolas'
        font.size = Pt(9)

# Append refined wording
add_h(1, 'Refinements and Code Examples')
add_p('This refinement clarifies core behaviours and provides compact code examples for reviewers and future maintainers.')

add_h(2, 'Secrets: concise description')
add_p('`get_secret(name, env_fallback=None)` resolves secrets with the following priority:')
add_p('- 1) Key Vault via `DefaultAzureCredential` when `KEY_VAULT_URL` is configured.')
add_p('- 2) Environment variable named by `env_fallback` when provided.')
add_p('- 3) Throws `RuntimeError` if unresolved to fail fast and avoid silent misconfiguration.')

add_h(2, 'Activity log: concise behaviour')
add_p('Activity logs are created for admin actions that modify tracked fields. Each log includes actor metadata, action, affected ticket, changed fields, and previous values.')

add_h(2, 'Notifications: concise behaviour')
add_p('Notifications are created per affected ticket and persist a `read` flag; clients query unread counts and mark items read via the API.')

add_h(2, 'Code snippets (representative)')
add_code_block([
    'def create_activity_log(actor, actor_type, action, ticket_id, updated_fields, old_values):',
    "    entry = {\n        'id': f'LOG-{int(time.time())}-{ticket_id}',\n        'actor': actor, 'actor_type': actor_type, 'action': action,\n        'ticket_id': ticket_id, 'timestamp': datetime.utcnow().isoformat(),\n        'updated_fields': updated_fields, 'old_values': old_values\n    }",
    "    write_json_blob('activitylogs.json', entry)  # append or upsert depending on store semantics",
])

add_code_block([
    'def create_notification(email, ticket_id, message, updated_fields=None):',
    "    notif = { 'id': f'NOTIF-{uuid.uuid4()}', 'email': email, 'ticket_id': ticket_id,",
    "              'message': message, 'updated_fields': updated_fields or [], 'timestamp': datetime.utcnow().isoformat(), 'read': False }",
    "    write_json_blob('notifications.json', notif)",
])

add_code_block([
    'def get_secret(secret_name, env_fallback=None):',
    "    if KEY_VAULT_URL:\n        try:\n            return keyvault_client.get_secret(secret_name)\n        except Exception:\n            pass",
    "    if env_fallback and os.getenv(env_fallback):\n        return os.getenv(env_fallback)",
    "    raise RuntimeError('Secret {} not found'.format(secret_name))",
])

# Save refined copy
new_path = src.with_name(src.stem + '_refined' + src.suffix)
try:
    doc.save(str(new_path))
    print('Refined write-up saved to', new_path)
except Exception as e:
    print('Failed to save refined write-up:', e)
    raise

# Basic validation
doc2 = Document(str(new_path))
found = any('Refinements and Code Examples' in p.text for p in doc2.paragraphs)
print('Validation:', 'found' if found else 'NOT found', 'heading in', new_path)
