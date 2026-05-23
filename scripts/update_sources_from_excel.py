from pathlib import Path
import shutil
from datetime import datetime
from openpyxl import load_workbook
from docx import Document

excel = Path(r"D:\Documents\MyMahir Docs\Course Work\Task Resources - Lee Xue Bao.xlsx")
docx = Path(r"D:\Documents\MyMahir Docs\Course Work\Project Submissions\QuickAid_WriteUp_updated.docx")

if not excel.exists():
    print("ERROR: Excel not found:", excel)
    raise SystemExit(2)
if not docx.exists():
    print("ERROR: Docx not found:", docx)
    raise SystemExit(2)

backup = docx.with_name(docx.stem + '_backup_sources_' + datetime.now().strftime('%Y%m%d_%H%M%S') + docx.suffix)
shutil.copy2(docx, backup)
print('Backup created:', backup)

wb = load_workbook(str(excel), data_only=True)
sheet = wb.active
sources = []
for row in sheet.iter_rows(min_row=1, values_only=True):
    # take first non-empty cell in the row as a source entry
    cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
    if cells:
        sources.append(cells[0])
# dedupe preserving order
seen = set()
uniq = []
for s in sources:
    if s not in seen:
        seen.add(s)
        uniq.append(s)
sources = uniq
print('Sources extracted:', len(sources))

# open docx and insert sources

doc = Document(str(docx))
found = None
for para in doc.paragraphs:
    if para.text.strip().lower() in ('sources', 'references', 'bibliography'):
        found = para
        break

if found is None:
    doc.add_paragraph('Sources')
    for s in sources:
        doc.add_paragraph(s)
else:
    # Insert a marker and then insert each source after the heading (preserve order)
    found.insert_paragraph_after('[Sources updated from Excel]')
    for s in reversed(sources):
        found.insert_paragraph_after(s)

# save
try:
    doc.save(str(docx))
    print('Docx updated:', docx)
except PermissionError as e:
    alt = docx.with_name(docx.stem + '_sources_updated' + docx.suffix)
    doc.save(str(alt))
    print('File was locked; saved updated copy to:', alt)
    print('Backup is at:', backup)
except Exception as e:
    print('ERROR saving docx:', e)
    print('Backup is at:', backup)
    raise
