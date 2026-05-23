from docx import Document
from datetime import datetime
from pathlib import Path
import shutil

source = Path(r"D:\Documents\MyMahir Docs\Course Work\Project Submissions\QuickAid_WriteUp_updated.docx")
if not source.exists():
    print('Source write-up not found:', source)
    raise SystemExit(1)

backup = source.with_name(source.stem + '_backup_' + datetime.now().strftime('%Y%m%d_%H%M%S') + source.suffix)
shutil.copy2(source, backup)
print('Backup created:', backup)

doc = Document(str(source))

replacements = {
    'No admin interface: There is currently no way for support staff to view all tickets or update ticket statuses. This would be the most valuable next feature to implement.':
        'Admin interface already exists: QuickAid includes a staff-facing admin console with an Overview dashboard, Manage Tickets table, bulk edit tools, and Access Requests handling. The more realistic improvement is richer reporting, exports, and automation for assignment workflows.',
    'No authentication: Any user can submit a ticket or retrieve tickets for any email address by guessing. A login system or email verification step would improve security in a real deployment.':
        'Authentication is already implemented for students and admins. Students sign in through the app login flow, while admins use Microsoft Entra ID via MSAL and backend token validation. A future improvement is stronger role provisioning, session hardening, and optional MFA enforcement.',
}

changed = 0
for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    if text in replacements:
        paragraph.text = replacements[text]
        changed += 1

# Add a short explicit Entra note near the end if the document does not already call it out clearly.
extra_note = (
    'Microsoft Entra ID is used in the admin sign-in path. The frontend is configured with the tenant, client ID, and API scope in `config.js`, '
    'and the backend verifies Entra bearer tokens before allowing admin-only operations such as ticket management and access request review.'
)
if not any('Microsoft Entra ID is used in the admin sign-in path' in p.text for p in doc.paragraphs):
    doc.add_heading('Microsoft Entra ID Integration', level=2)
    doc.add_paragraph(extra_note)

output = source.with_name(source.stem + '_revised.docx')
doc.save(str(output))
print('Revised write-up saved to', output)
print('Replacements applied:', changed)

# Validate key text is present
check_doc = Document(str(output))
checks = [
    'Admin interface already exists',
    'Authentication is already implemented for students and admins',
    'Microsoft Entra ID is used in the admin sign-in path',
]
for needle in checks:
    print('CHECK', needle, '=>', any(needle in p.text for p in check_doc.paragraphs))
